"""Global pytest fixtures and configuration for myAudible tests."""

import os
from pathlib import Path
import tempfile
import pytest


@pytest.fixture
def tmp_dir():
    """Create a temporary directory for test file operations."""
    with tempfile.TemporaryDirectory() as tmp_path:
        yield Path(tmp_path)


@pytest.fixture
def sample_text():
    """Sample text for testing text extraction and chunking."""
    return """
This is a sample paragraph for testing purposes. It contains multiple sentences that should be split into chunks.

The second paragraph continues the text with additional content for semantic chunking. This helps verify that the chunking algorithm works correctly.
    """


@pytest.fixture
def sample_markdown_text():
    """Sample markdown text for testing markdown extraction."""
    return """
# Sample Document

This is a **bold** statement and this is *italic* text.

## Section 1

This is a code block:

```python
def hello():
    print("Hello, World!")
```

## Section 2

This is a regular paragraph with some more content.
    """


@pytest.fixture
def mock_config(tmp_dir):
    """Create a mock configuration object for testing."""
    from myaudible.config import AppConfig, TTSConfig, ProcessingConfig, LoggingConfig
    
    config = AppConfig(
        input_dir=tmp_dir / "input",
        output_dir=tmp_dir / "output",
        tts=TTSConfig(
            endpoint="http://localhost:8008/v1/audio/speech",
            retry_attempts=3,
        ),
        processing=ProcessingConfig(
            chunk_size=500,
            max_parallel=2,
        ),
        logging=LoggingConfig(
            level="INFO",
        ),
    )
    return config


@pytest.fixture
def sample_file(tmp_dir):
    """Helper to create a temporary test file."""
    def _create_file(suffix=".txt", content=""):
        file_path = tmp_dir / f"sample{suffix}"
        file_path.write_text(content, encoding="utf-8")
        return file_path
    
    return _create_file


@pytest.fixture
def sample_pdf_file(tmp_dir):
    """Helper to create a sample PDF file for testing."""
    def _create_pdf():
        import os
        # Create a minimal PDF for testing
        # In real tests, this would use a real PDF file
        pdf_path = tmp_dir / "sample.pdf"
        # For now, create an empty file - actual PDF tests would use real PDFs
        pdf_path.touch()
        return pdf_path
    
    return _create_pdf


@pytest.fixture
def sample_docx_file(tmp_dir):
    """Helper to create a sample DOCX file for testing."""
    def _create_docx():
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("This is a sample Word document paragraph.")
        doc.add_paragraph("This is another paragraph for testing.")
        
        docx_path = tmp_dir / "sample.docx"
        doc.save(docx_path)
        return docx_path
    
    return _create_docx
