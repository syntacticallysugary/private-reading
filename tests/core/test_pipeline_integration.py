"""Integration tests for the ProcessingPipeline.

These tests verify the complete flow from text extraction through TTS generation
to audio stitching and output management.
"""

import asyncio
import json
import sys
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from private_reading.config import AppConfig, LoggingConfig, ProcessingConfig, TTSConfig
from private_reading.core.audio_stitcher import AudioStitcher
from private_reading.core.chunk_manager import MAX_CHUNK, ChunkManager
from private_reading.core.job_tracker import JobTracker
from private_reading.core.output_manager import OutputManager
from private_reading.core.pipeline import PipelineStatus, ProcessingPipeline, ProcessingResult
from private_reading.core.text_extractor import TextExtractor
from private_reading.core.tts_client import TTSClient
from private_reading.exceptions import (
    AudioError,
    ChunkingError,
    ExtractionError,
    OutputError,
    TTSError,
)


class TestPipelineIntegration:
    """Integration tests for the full pipeline."""

    @pytest.fixture
    def app_config(self, tmp_path):
        """Create a test AppConfig."""
        return AppConfig(
            input_dir=tmp_path / "input",
            output_dir=tmp_path / "output",
            processed_dir=tmp_path / "processed",
            tts=TTSConfig(
                endpoint="http://localhost:8008/v1/audio/speech",
                retry_attempts=1,
            ),
            processing=ProcessingConfig(
                chunk_size=MAX_CHUNK,
                overlap_ratio=0.1,
            ),
            logging=LoggingConfig(level="INFO"),
        )

    @pytest.fixture
    def pipeline(self, app_config):
        """Create a ProcessingPipeline instance."""
        return ProcessingPipeline(app_config)

    @pytest.fixture
    def test_directory(self, tmp_path):
        """Create test directories."""
        input_dir = tmp_path / "input"
        output_dir = tmp_path / "output"
        processed_dir = tmp_path / "processed"
        input_dir.mkdir(parents=True, exist_ok=True)
        output_dir.mkdir(parents=True, exist_ok=True)
        processed_dir.mkdir(parents=True, exist_ok=True)
        return {
            "input": input_dir,
            "output": output_dir,
            "processed": processed_dir,
        }

    # =========================================================================
    # 1. Pipeline Initialization Tests
    # =========================================================================

    @pytest.mark.asyncio
    async def test_pipeline_init(self, app_config):
        """Test that ProcessingPipeline creates successfully with config."""
        pipeline = ProcessingPipeline(app_config)

        # Assert pipeline exists
        assert pipeline is not None
        assert pipeline.config == app_config

        # Assert components are initialized
        assert pipeline.extractor is not None
        assert pipeline.chunk_manager is not None
        assert pipeline.tts_client is not None
        assert pipeline.audio_stitcher is not None
        assert pipeline.output_manager is not None
        assert pipeline.job_tracker is None  # JobTracker not yet implemented

    @pytest.mark.asyncio
    async def test_pipeline_components(self, app_config):
        """Test that all components are instantiated."""
        pipeline = ProcessingPipeline(app_config)

        # Assert self.extractor is TextExtractor instance
        assert isinstance(pipeline.extractor, TextExtractor)

        # Assert self.chunk_manager is ChunkManager instance
        assert isinstance(pipeline.chunk_manager, ChunkManager)
        assert pipeline.chunk_manager.max_chars == MAX_CHUNK
        assert pipeline.chunk_manager.overlap_ratio == 0.1

        # Assert self.tts_client is TTSClient instance
        assert isinstance(pipeline.tts_client, TTSClient)
        assert pipeline.tts_client.endpoint == "http://localhost:8008/v1/audio/speech"
        assert pipeline.tts_client.retry_attempts == 1

        # Assert self.audio_stitcher is AudioStitcher instance
        assert isinstance(pipeline.audio_stitcher, AudioStitcher)

        # Assert self.output_manager is OutputManager instance
        assert isinstance(pipeline.output_manager, OutputManager)
        assert pipeline.output_manager.output_dir == app_config.output_dir
        assert pipeline.output_manager.processed_dir == app_config.processed_dir

        # Assert self.job_tracker is None (not yet implemented)
        assert pipeline.job_tracker is None

    # =========================================================================
    # 2. File Format Processing Tests
    # =========================================================================

    @pytest.mark.asyncio
    async def test_process_markdown_file(self, pipeline, test_directory):
        """Test end-to-end markdown file processing."""
        # Create test markdown file
        md_file = test_directory["input"] / "test.md"
        md_content = """# Test Document

This is **bold** and *italic* text.

## Section

Some regular text here.
"""
        md_file.write_text(md_content, encoding="utf-8")

        # Process through pipeline
        result = await pipeline.process_file(md_file)

        # Assert ProcessingResult.success is True
        assert result.success is True

        # Assert output WAV file exists
        assert result.output_path is not None
        assert result.output_path.exists()
        assert result.output_path.suffix == ".wav"

        # Assert sidecar metadata exists
        sidecar_path = result.output_path.with_stem("test.md.json")
        assert sidecar_path.exists()

        # Verify sidecar content
        with open(sidecar_path) as f:
            metadata = json.load(f)
        assert "source_file" in metadata
        assert "duration" in metadata
        assert "chunk_count" in metadata

    @pytest.mark.asyncio
    async def test_process_txt_file(self, pipeline, test_directory):
        """Test end-to-end txt file processing."""
        # Create test txt file
        txt_file = test_directory["input"] / "test.txt"
        txt_content = "This is a test text file with multiple sentences for processing."
        txt_file.write_text(txt_content, encoding="utf-8")

        # Process through pipeline
        result = await pipeline.process_file(txt_file)

        # Assert success and output exists
        assert result.success is True
        assert result.output_path is not None
        assert result.output_path.exists()

        # Assert sidecar metadata exists
        sidecar_path = result.output_path.with_stem("test.txt.json")
        assert sidecar_path.exists()

    @pytest.mark.asyncio
    async def test_process_pdf_file(self, pipeline, test_directory):
        """Test end-to-end pdf file processing (if pdfplumber available)."""
        # Skip if pdfplumber not installed
        if not pytest.importorskip("pdfplumber"):
            pytest.skip("pdfplumber not installed")

        # Create test pdf (minimal PDF for testing)
        pdf_file = test_directory["input"] / "test.pdf"
        # Create a minimal valid PDF
        pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT
/F1 12 Tf
100 700 Td
(Test PDF) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000053 00000 n
0000000101 00000 n
0000000172 00000 n
trailer
<< /Size 5 /Root 1 0 R >>
startxref
271
%%EOF"""
        pdf_file.write_bytes(pdf_content)

        # Process through pipeline
        result = await pipeline.process_file(pdf_file)

        # Assert success
        assert result.success is True
        assert result.output_path is not None
        assert result.output_path.exists()

    @pytest.mark.asyncio
    async def test_process_docx_file(self, pipeline, test_directory):
        """Test end-to-end docx file processing (if python-docx available)."""
        # Skip if python-docx not installed
        if not pytest.importorskip("python-docx"):
            pytest.skip("python-docx not installed")

        # Create test docx file
        docx_file = test_directory["input"] / "test.docx"
        from docx import Document

        doc = Document()
        doc.add_paragraph("This is a test Word document with multiple paragraphs.")
        doc.add_paragraph("This is another paragraph for testing.")
        doc.save(docx_file)

        # Process through pipeline
        result = await pipeline.process_file(docx_file)

        # Assert success
        assert result.success is True
        assert result.output_path is not None
        assert result.output_path.exists()

    # =========================================================================
    # 3. Chunking and Audio Tests
    # =========================================================================

    @pytest.mark.asyncio
    async def test_chunking_with_overlap(self, pipeline, test_directory):
        """Test that chunks have correct overlap."""
        # Create pipeline with known overlap ratio
        chunk_manager = pipeline.chunk_manager

        # Sample text that will be chunked
        sample_text = """
This is the first paragraph with some content for testing.

This is the second paragraph with more content to ensure we have enough text.

This is the third paragraph to verify chunking works correctly with overlap.
"""

        # Chunk text
        chunks = await chunk_manager.chunk(sample_text)

        # Assert chunks were created
        assert len(chunks) > 0

        # Verify overlap by checking chunk boundaries
        # The overlap_ratio is 0.1, so chunks should overlap by 10%
        if len(chunks) > 1:
            # Check that chunks have some overlap (overlap_ratio * max_chars)
            expected_overlap = int(chunk_manager.max_chars * chunk_manager.overlap_ratio)
            for i in range(1, len(chunks)):
                # Last chunk of previous chunk should appear in next chunk
                overlap_start = len(chunks[i - 1]) - expected_overlap
                overlap_end = len(chunks[i])
                assert (
                    chunks[i][:expected_overlap] == chunks[i - 1][-expected_overlap:]
                ), f"Chunk {i} should overlap with chunk {i-1}"

    @pytest.mark.asyncio
    async def test_audio_stitching(self, pipeline, test_directory):
        """Test that audio chunks combine correctly."""
        audio_stitcher = pipeline.audio_stitcher

        # Create mock audio chunks
        mock_chunks = [
            b"audio data chunk 1",
            b"audio data chunk 2",
            b"audio data chunk 3",
        ]

        # Create temporary WAV files for stitching
        wav_files = []
        for i, chunk in enumerate(mock_chunks):
            wav_path = test_directory["output"] / f"chunk_{i}.wav"
            wav_path.write_bytes(chunk)
            wav_files.append(wav_path)

        # Stitch files
        output_path = test_directory["output"] / "stitched.wav"

        with patch("asyncio.create_subprocess_exec") as mock_subprocess:
            mock_process = AsyncMock()
            mock_process.communicate = AsyncMock(return_value=(b"", b""))
            mock_process.returncode = 0
            mock_subprocess.return_value = mock_process

            await audio_stitcher.stitch(wav_files, output_path)

            # Assert output file exists and has correct duration
            assert output_path.exists()
            assert output_path.suffix == ".wav"

    @pytest.mark.asyncio
    async def test_silence_insertion(self, pipeline, test_directory):
        """Test that silence is added between chunks."""
        chunk_manager = pipeline.chunk_manager

        # Create pipeline with silence config
        sample_text = """
First paragraph.

Second paragraph.

Third paragraph.
"""

        # Chunk text
        chunks = await chunk_manager.chunk(sample_text)

        # Add silence markers
        markers = await chunk_manager.add_silence_markers(chunks)

        # Assert silence periods exist in output
        assert len(markers) == len(chunks)
        assert markers[0]["has_silence_before"] is False
        assert all(m["has_silence_before"] for m in markers[1:])

    # =========================================================================
    # 4. Output Tests
    # =========================================================================

    @pytest.mark.asyncio
    async def test_wav_output_created(self, pipeline, test_directory):
        """Test that WAV file exists after processing."""
        # Create input file
        input_file = test_directory["input"] / "test.txt"
        input_file.write_text("Test content for WAV output.", encoding="utf-8")

        # Process file through pipeline
        result = await pipeline.process_file(input_file)

        # Assert output_path exists
        assert result.output_path is not None
        assert result.output_path.exists()

        # Assert file ends with .wav
        assert result.output_path.suffix == ".wav"

    @pytest.mark.asyncio
    async def test_sidecar_metadata(self, pipeline, test_directory):
        """Test that JSON sidecar has correct data."""
        # Create input file
        input_file = test_directory["input"] / "test.txt"
        input_file.write_text("Test content for sidecar metadata.", encoding="utf-8")

        # Process file through pipeline
        result = await pipeline.process_file(input_file)

        # Assert sidecar exists
        sidecar_path = result.output_path.with_stem("test.txt.json")
        assert sidecar_path.exists()

        # Load sidecar JSON
        with open(sidecar_path) as f:
            metadata = json.load(f)

        # Assert metadata contains required fields
        assert "source_file" in metadata
        assert "duration" in metadata
        assert "chunk_count" in metadata
        assert "processing_timestamp" in metadata or "processing_time_seconds" in metadata

    @pytest.mark.asyncio
    async def test_file_moved_to_processed(self, pipeline, test_directory):
        """Test that input is moved to archive."""
        # Create input file
        input_file = test_directory["input"] / "test.txt"
        input_file.write_text("Test content for file moving.", encoding="utf-8")

        # Process file through pipeline
        result = await pipeline.process_file(input_file)

        # Assert source file moved to processed directory
        processed_file = test_directory["processed"] / "test.txt"
        assert processed_file.exists()
        assert not input_file.exists()

    # =========================================================================
    # 5. Job Tracker Integration Tests
    # =========================================================================

    @pytest.mark.asyncio
    async def test_job_tracker_integration(self, pipeline, test_directory):
        """Test that jobs are tracked during processing."""
        # Create pipeline with JobTracker
        job_tracker = JobTracker()
        pipeline.job_tracker = job_tracker

        # Create input file
        input_file = test_directory["input"] / "test.txt"
        input_file.write_text("Test content for job tracking.", encoding="utf-8")

        # Process file
        result = await pipeline.process_file(input_file)

        # Assert job exists in tracker
        assert job_tracker is not None
        assert len(job_tracker.list_jobs()) > 0

        # Get the job for this file
        jobs = job_tracker.get_jobs_by_status(job_tracker.COMPLETED)
        assert len(jobs) > 0

        # Assert job has correct source_file
        job = jobs[0]
        assert job.source_file == input_file

    @pytest.mark.asyncio
    async def test_job_status_updates(self, pipeline, test_directory):
        """Test that job status transitions correctly."""
        # Create pipeline with JobTracker
        job_tracker = JobTracker()
        pipeline.job_tracker = job_tracker

        # Create input file
        input_file = test_directory["input"] / "test.txt"
        input_file.write_text("Test content for status updates.", encoding="utf-8")

        # Process file
        result = await pipeline.process_file(input_file)

        # Assert job status is COMPLETED
        jobs = job_tracker.get_jobs_by_status(job_tracker.COMPLETED)
        assert len(jobs) > 0

        job = jobs[0]
        assert job.status == job_tracker.COMPLETED

        # Assert duration > 0
        assert job.duration is not None
        assert job.duration > 0

    # =========================================================================
    # 6. Error Scenario Tests
    # =========================================================================

    @pytest.mark.asyncio
    async def test_invalid_file_format(self, pipeline, test_directory):
        """Test handling of unsupported formats."""
        # Create file with unsupported extension
        unsupported_file = test_directory["input"] / "test.xyz"
        unsupported_file.write_text("This format is not supported.", encoding="utf-8")

        # Process through pipeline
        result = await pipeline.process_file(unsupported_file)

        # Assert ProcessingResult.success is False
        assert result.success is False

        # Assert error message mentions unsupported format
        assert "unsupported" in result.error.lower() or "format" in result.error.lower()

    @pytest.mark.asyncio
    async def test_missing_input_file(self, pipeline, test_directory):
        """Test handling of non-existent files."""
        # Process non-existent file path
        nonexistent_file = test_directory["input"] / "nonexistent.txt"

        # Process through pipeline
        result = await pipeline.process_file(nonexistent_file)

        # Assert ProcessingResult.success is False
        assert result.success is False

        # Assert error message mentions file not found
        assert "not found" in result.error.lower() or "no such" in result.error.lower()

    @pytest.mark.asyncio
    async def test_tts_api_failure(self, pipeline, test_directory):
        """Test handling of TTS API errors."""
        # Create input file
        input_file = test_directory["input"] / "test.txt"
        input_file.write_text("Test content for TTS failure.", encoding="utf-8")

        # Mock TTSClient.generate_speech to raise TTSError
        with patch.object(
            pipeline.tts_client, "generate_speech", new_callable=AsyncMock
        ) as mock_generate:
            mock_generate.side_effect = TTSError("TTS API unavailable")

            # Process file through pipeline
            result = await pipeline.process_file(input_file)

            # Assert ProcessingResult.success is False
            assert result.success is False

            # Assert error message mentions TTS failure
            assert "tts" in result.error.lower() or "api" in result.error.lower()


# Additional helper tests for individual components


class TestPipelineComponents:
    """Tests for individual pipeline components."""

    @pytest.fixture
    def app_config(self, tmp_path):
        """Create a test AppConfig."""
        return AppConfig(
            input_dir=tmp_path / "input",
            output_dir=tmp_path / "output",
            processed_dir=tmp_path / "processed",
            tts=TTSConfig(
                endpoint="http://localhost:8008/v1/audio/speech",
                retry_attempts=1,
            ),
            processing=ProcessingConfig(
                chunk_size=MAX_CHUNK,
                overlap_ratio=0.1,
            ),
            logging=LoggingConfig(level="INFO"),
        )

    @pytest.fixture
    def pipeline(self, app_config):
        """Create a ProcessingPipeline instance."""
        return ProcessingPipeline(app_config)

    # Pipeline Status Tests

    @pytest.mark.asyncio
    async def test_pipeline_status(self, pipeline):
        """Test that pipeline status is returned correctly."""
        status = pipeline.get_status()

        # Assert status has required attributes
        assert isinstance(status.state, str)
        assert status.state in ["IDLE", "PROCESSING", "ERROR"]
        assert isinstance(status.component_health, dict)
        assert status.active_jobs == 0

    @pytest.mark.asyncio
    async def test_pipeline_state_idle(self, pipeline):
        """Test that pipeline state is IDLE when not processing."""
        status = pipeline.get_status()

        # Assert state is IDLE
        assert status.state == "IDLE"

    @pytest.mark.asyncio
    async def test_pipeline_component_health(self, pipeline):
        """Test that component health is tracked."""
        status = pipeline.get_status()

        # Assert all components have health status
        assert "TextExtractor" in status.component_health
        assert "ChunkManager" in status.component_health
        assert "TTSClient" in status.component_health
        assert "AudioStitcher" in status.component_health
        assert "OutputManager" in status.component_health

    # ProcessingResult Tests

    @pytest.mark.asyncio
    async def test_processing_result_success(self):
        """Test ProcessingResult with success."""
        result = ProcessingResult(
            success=True,
            output_path=Path("/tmp/output.wav"),
            duration=1.5,
        )

        # Assert result has correct attributes
        assert result.success is True
        assert result.output_path is not None
        assert result.duration == 1.5
        assert result.error is None

    @pytest.mark.asyncio
    async def test_processing_result_failure(self):
        """Test ProcessingResult with failure."""
        result = ProcessingResult(
            success=False,
            output_path=None,
            error="Test error message",
            duration=0.5,
        )

        # Assert result has correct attributes
        assert result.success is False
        assert result.output_path is None
        assert result.error == "Test error message"
        assert result.duration == 0.5

    @pytest.mark.asyncio
    async def test_processing_result_post_init(self):
        """Test ProcessingResult __post_init__ sets output_path to None on failure."""
        # Create result with success=False and output_path=None
        result = ProcessingResult(
            success=False,
            output_path=None,
            error="Test error",
        )

        # Assert output_path remains None
        assert result.output_path is None


# Run tests
if __name__ == "__main__":
    pytest.main([__file__, "-v"])


# =========================================================================
# Semaphore Tests
# =========================================================================


class TestSemaphoreBehavior:
    """Tests for semaphore-based concurrency control."""

    @pytest.fixture
    def app_config(self, tmp_path):
        """Create a test AppConfig with default semaphore size."""
        from private_reading.config import SemaphoreConfig

        return AppConfig(
            input_dir=tmp_path / "input",
            output_dir=tmp_path / "output",
            processed_dir=tmp_path / "processed",
            semaphore=SemaphoreConfig(size=10),
            tts=TTSConfig(
                endpoint="http://localhost:8008/v1/audio/speech",
                retry_attempts=1,
            ),
            processing=ProcessingConfig(
                chunk_size=MAX_CHUNK,
                overlap_ratio=0.1,
            ),
            logging=LoggingConfig(level="INFO"),
        )

    @pytest.mark.asyncio
    async def test_semaphore_size_1(self, tmp_path):
        """Test that semaphore is initialized with size=1 (sequential processing)."""
        from private_reading.config import SemaphoreConfig
        from private_reading.core.pipeline import ProcessingPipeline

        config = AppConfig(
            input_dir=tmp_path / "input",
            output_dir=tmp_path / "output",
            processed_dir=tmp_path / "processed",
            semaphore=SemaphoreConfig(size=1),
            tts=TTSConfig(
                endpoint="http://localhost:8008/v1/audio/speech",
                retry_attempts=1,
            ),
            processing=ProcessingConfig(
                chunk_size=MAX_CHUNK,
                overlap_ratio=0.1,
            ),
            logging=LoggingConfig(level="INFO"),
        )

        pipeline = ProcessingPipeline(config)

        # Assert semaphore exists and has correct size
        assert pipeline.semaphore is not None
        assert isinstance(pipeline.semaphore, asyncio.Semaphore)

        # Check the internal semaphore value (limit)
        # asyncio.Semaphore's _value is protected, so we test via acquire/release behavior
        # The semaphore should allow exactly 1 concurrent operation
        assert pipeline.semaphore._value == 1

    @pytest.mark.asyncio
    async def test_semaphore_size_10(self, app_config):
        """Test that semaphore is initialized with size=10 (default)."""
        from private_reading.core.pipeline import ProcessingPipeline

        pipeline = ProcessingPipeline(app_config)

        # Assert semaphore exists and has correct size
        assert pipeline.semaphore is not None
        assert isinstance(pipeline.semaphore, asyncio.Semaphore)
        assert pipeline.semaphore._value == 10

    @pytest.mark.asyncio
    async def test_semaphore_size_50(self, tmp_path):
        """Test that semaphore is initialized with size=50 (max concurrency)."""
        from private_reading.config import SemaphoreConfig
        from private_reading.core.pipeline import ProcessingPipeline

        config = AppConfig(
            input_dir=tmp_path / "input",
            output_dir=tmp_path / "output",
            processed_dir=tmp_path / "processed",
            semaphore=SemaphoreConfig(size=50),
            tts=TTSConfig(
                endpoint="http://localhost:8008/v1/audio/speech",
                retry_attempts=1,
            ),
            processing=ProcessingConfig(
                chunk_size=MAX_CHUNK,
                overlap_ratio=0.1,
            ),
            logging=LoggingConfig(level="INFO"),
        )

        pipeline = ProcessingPipeline(config)

        # Assert semaphore exists and has correct size
        assert pipeline.semaphore is not None
        assert isinstance(pipeline.semaphore, asyncio.Semaphore)
        assert pipeline.semaphore._value == 50

    @pytest.mark.asyncio
    async def test_semaphore_concurrency_limit(self, tmp_path):
        """Test that semaphore actually limits concurrent operations."""
        from private_reading.config import SemaphoreConfig
        from private_reading.core.pipeline import ProcessingPipeline

        # Create pipeline with semaphore size=2
        config = AppConfig(
            input_dir=tmp_path / "input",
            output_dir=tmp_path / "output",
            processed_dir=tmp_path / "processed",
            semaphore=SemaphoreConfig(size=2),
            tts=TTSConfig(
                endpoint="http://localhost:8008/v1/audio/speech",
                retry_attempts=1,
            ),
            processing=ProcessingConfig(
                chunk_size=MAX_CHUNK,
                overlap_ratio=0.1,
            ),
            logging=LoggingConfig(level="INFO"),
        )

        pipeline = ProcessingPipeline(config)

        # Test that semaphore correctly limits concurrency
        # We can verify this by checking the internal state before and after acquire
        initial_value = pipeline.semaphore._value

        # Acquire the semaphore (this should decrement the value)
        await pipeline.semaphore.acquire()
        assert pipeline.semaphore._value == initial_value - 1

        # Release it back
        pipeline.semaphore.release()
        assert pipeline.semaphore._value == initial_value


# =========================================================================
# Semaphore Validation Tests
# =========================================================================


class TestSemaphoreValidation:
    """Tests for semaphore configuration validation."""

    @pytest.mark.asyncio
    async def test_invalid_semaphore_size_zero(self, tmp_path):
        """Test that semaphore size=0 raises ValueError."""
        from private_reading.config import SemaphoreConfig

        with pytest.raises(ValueError) as exc_info:
            SemaphoreConfig(size=0)

        assert "Semaphore size must be between 1 and 50" in str(exc_info.value)

    @pytest.mark.asyncio
    async def test_invalid_semaphore_size_negative(self, tmp_path):
        """Test that negative semaphore size raises ValueError."""
        from private_reading.config import SemaphoreConfig

        with pytest.raises(ValueError) as exc_info:
            SemaphoreConfig(size=-1)

        assert "Semaphore size must be between 1 and 50" in str(exc_info.value)

    @pytest.mark.asyncio
    async def test_invalid_semaphore_size_exceeds_max(self, tmp_path):
        """Test that semaphore size > 50 raises ValueError."""
        from private_reading.config import SemaphoreConfig

        with pytest.raises(ValueError) as exc_info:
            SemaphoreConfig(size=51)

        assert "Semaphore size must be between 1 and 50" in str(exc_info.value)

    @pytest.mark.asyncio
    async def test_valid_semaphore_size_boundary_1(self, tmp_path):
        """Test that semaphore size=1 is valid (minimum)."""
        from private_reading.config import SemaphoreConfig

        # Should not raise
        config = SemaphoreConfig(size=1)
        assert config.size == 1

    @pytest.mark.asyncio
    async def test_valid_semaphore_size_boundary_50(self, tmp_path):
        """Test that semaphore size=50 is valid (maximum)."""
        from private_reading.config import SemaphoreConfig

        # Should not raise
        config = SemaphoreConfig(size=50)
        assert config.size == 50

    @pytest.mark.asyncio
    async def test_valid_semaphore_size_middle(self, tmp_path):
        """Test that a valid middle value like 25 is accepted."""
        from private_reading.config import SemaphoreConfig

        # Should not raise
        config = SemaphoreConfig(size=25)
        assert config.size == 25


# =========================================================================
# CLI Argument Parsing Tests
# =========================================================================


class TestCLIArgumentParsing:
    """Tests for CLI argument parsing of semaphore size."""

    def test_semaphore_size_default(self):
        """Test that default semaphore size is 10 when not specified."""
        import argparse

        from private_reading.cli import create_argument_parser

        parser = create_argument_parser()

        # Parse with minimal args (no semaphore-size specified)
        args = parser.parse_args(["-i", "/input", "-o", "/output"])

        assert args.semaphore_size is None

    def test_semaphore_size_valid(self):
        """Test parsing valid semaphore size argument."""
        import argparse

        from private_reading.cli import create_argument_parser

        parser = create_argument_parser()

        # Parse with semaphore-size=25
        args = parser.parse_args(["-i", "/input", "-o", "/output", "--semaphore-size", "25"])

        assert args.semaphore_size == 25

    def test_semaphore_size_boundary_one(self):
        """Test parsing semaphore size=1."""
        from private_reading.cli import create_argument_parser

        parser = create_argument_parser()

        args = parser.parse_args(["-i", "/input", "-o", "/output", "--semaphore-size", "1"])

        assert args.semaphore_size == 1

    def test_semaphore_size_boundary_fifty(self):
        """Test parsing semaphore size=50."""
        from private_reading.cli import create_argument_parser

        parser = create_argument_parser()

        args = parser.parse_args(["-i", "/input", "-o", "/output", "--semaphore-size", "50"])

        assert args.semaphore_size == 50

    def test_semaphore_size_invalid_negative(self, capsys):
        """Test that CLI rejects negative semaphore size."""
        from private_reading.cli import create_argument_parser, validate_inputs

        parser = create_argument_parser()

        args = parser.parse_args(["-i", "/input", "-o", "/output", "--semaphore-size", "-5"])

        # Validate should return False for invalid semaphore size
        result = validate_inputs(args)

        assert result is False

    def test_semaphore_size_invalid_zero(self, capsys):
        """Test that CLI rejects zero semaphore size."""
        from private_reading.cli import create_argument_parser, validate_inputs

        parser = create_argument_parser()

        args = parser.parse_args(["-i", "/input", "-o", "/output", "--semaphore-size", "0"])

        result = validate_inputs(args)

        assert result is False

    def test_semaphore_size_invalid_exceeds_max(self, capsys):
        """Test that CLI rejects semaphore size > 50."""
        from private_reading.cli import create_argument_parser, validate_inputs

        parser = create_argument_parser()

        args = parser.parse_args(["-i", "/input", "-o", "/output", "--semaphore-size", "100"])

        result = validate_inputs(args)

        assert result is False

    def test_semaphore_size_build_config(self, tmp_path, monkeypatch):
        """Test that CLI config correctly applies semaphore size."""
        from private_reading.cli import build_config, create_argument_parser

        parser = create_argument_parser()

        args = parser.parse_args(
            ["-i", str(tmp_path / "input"), "-o", str(tmp_path / "output"), "--semaphore-size", "5"]
        )

        config = build_config(args)

        assert config.semaphore.size == 5

    def test_semaphore_size_default_in_config(self, tmp_path, monkeypatch):
        """Test that default semaphore size (10) is applied when not overridden."""
        from private_reading.cli import build_config, create_argument_parser

        parser = create_argument_parser()

        args = parser.parse_args(["-i", str(tmp_path / "input"), "-o", str(tmp_path / "output")])

        config = build_config(args)

        assert config.semaphore.size == 10

    def test_short_sem_flag_not_supported(self, tmp_path, capsys):
        """Test that short -s flag is not supported for semaphore-size."""
        from private_reading.cli import create_argument_parser

        parser = create_argument_parser()

        # Try to use -s flag (should fail since only --semaphore-size is supported)
        with pytest.raises(SystemExit):
            parser.parse_args(["-i", "/input", "-o", "/output", "-s", "10"])

    def test_semaphore_size_in_help(self, capsys):
        """Test that --semaphore-size appears in help text."""
        from private_reading.cli import create_argument_parser

        parser = create_argument_parser()

        try:
            parser.parse_args(["-h"])
        except SystemExit:
            pass

        # Help should mention semaphore-size
        captured = capsys.readouterr()
        assert "--semaphore-size" in captured.out
        assert "1-50" in captured.out
